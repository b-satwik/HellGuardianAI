import sys
import re

try:
    import docx
except ImportError:
    print("Dependencies missing. Please install python-docx:")
    print("pip install python-docx")
    sys.exit(1)

def convert_md_to_docx():
    doc = docx.Document()
    
    # Adjust default styles if desired
    # Set document title/header
    doc.add_heading("HellGuardianAI - Hackathon Submission", level=0)

    with open("hackathon_submission.md", "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        cleaned = line.strip()
        
        # Skip empty lines or markdown horizontal rules
        if not cleaned or cleaned == "---":
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            continue

        # Header 1
        if cleaned.startswith("# "):
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            doc.add_heading(cleaned[2:], level=1)
            
        # Header 2
        elif cleaned.startswith("## "):
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            doc.add_heading(cleaned[3:], level=2)
            
        # Header 3
        elif cleaned.startswith("### "):
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            doc.add_heading(cleaned[4:], level=3)
            
        # Unordered Lists
        elif cleaned.startswith("* ") or cleaned.startswith("- "):
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            
            # Clean bold markers inside list items
            item_text = cleaned[2:]
            item_text = re.sub(r"\*\*(.*?)\*\*", r"\1", item_text)
            doc.add_paragraph(item_text, style="List Bullet")
            
        # Table Parsing
        elif cleaned.startswith("|"):
            # Skip header separator lines (e.g., | :--- | :--- |)
            if "---" in cleaned or ":---" in cleaned:
                continue
            in_table = True
            cells = [c.strip() for c in cleaned.split("|")[1:-1]]
            table_data.append(cells)
            
        # Standard Paragraphs / Mermaid Diagrams / Architecture Flow
        else:
            if in_table and table_data:
                build_table(doc, table_data)
                table_data = []
                in_table = False
            
            # Skip code blocks / mermaid wrapper lines
            if cleaned.startswith("```") or cleaned == "```mermaid" or cleaned == "▼" or cleaned == "│":
                # Just add raw lines for text flow inside architecture diagrams
                if cleaned != "```" and cleaned != "```mermaid":
                    doc.add_paragraph(cleaned)
                continue
            
            # Clean bold markdown markers
            p_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
            doc.add_paragraph(p_text)

    # Flush any remaining table at the end of the file
    if in_table and table_data:
        build_table(doc, table_data)

    output_filename = "hackathon_submission.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename} in the project folder!")

def build_table(doc, table_data):
    if not table_data:
        return
    
    rows = len(table_data)
    cols = len(table_data[0])
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    
    for r_idx, row in enumerate(table_data):
        # Prevent index out of bounds in case of malformed markdown tables
        for c_idx, val in enumerate(row[:cols]):
            cell = table.cell(r_idx, c_idx)
            cell.text = val

if __name__ == "__main__":
    convert_md_to_docx()
